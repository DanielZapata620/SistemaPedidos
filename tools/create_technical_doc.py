from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = r"C:\Users\danie\OneDrive\Datos adjuntos\Documentos\ProyectoPery\Explicacion_Tecnica_Plataforma_Pedidos.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(107, 62, 29)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(7)
styles["Normal"].paragraph_format.line_spacing = 1.12

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Explicacion Tecnica del Proyecto")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(107, 62, 29)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Plataforma web de pedidos para negocio con cliente, administrador y sucursales")
run.font.name = "Arial"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(90, 84, 78)

doc.add_paragraph(
    "El sistema resuelve la necesidad de un negocio pequeno que requiere administrar productos, "
    "recibir pedidos de clientes y controlar el estado de cada pedido hasta que este listo para recoger. "
    "La aplicacion se desarrollo con separacion entre frontend, servicios y API, por lo que queda preparada "
    "para evolucionar posteriormente a una PWA."
)

add_heading(doc, "1. Proposito del sistema")
add_bullet(doc, "Permitir que el cliente consulte productos, arme un pedido, seleccione sucursal y confirme la compra para recoger en tienda.")
add_bullet(doc, "Permitir que el administrador gestione productos, sucursales y pedidos desde un panel de negocio.")
add_bullet(doc, "Guardar la informacion en MySQL para que productos, usuarios, sucursales y pedidos persistan aunque se cierre la aplicacion.")
add_bullet(doc, "Mostrar dashboard dinamico con productos activos, pedidos totales, pendientes y ventas acumuladas.")

add_heading(doc, "2. Tecnologias utilizadas")
tech_table = doc.add_table(rows=1, cols=2)
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tech_table.style = "Table Grid"
headers = tech_table.rows[0].cells
set_cell_text(headers[0], "Capa", True)
set_cell_text(headers[1], "Tecnologia / uso", True)
for c in headers:
    shade_cell(c, "F5E6CA")

rows = [
    ("Frontend", "Angular con componentes, rutas, formularios, servicios y consumo asincrono mediante fetch centralizado."),
    ("Backend", "ASP.NET Core API con controladores MVC, servicios, repositorio generico y DTOs con clases normales."),
    ("Base de datos", "MySQL con Entity Framework Core y creacion automatica de estructura al ejecutar la API."),
    ("Validaciones", "FluentValidation en la API y validaciones de formularios en Angular."),
    ("Mapeo", "AutoMapper con profiles para convertir entidades a DTOs."),
    ("Autenticacion", "JWT para sesion local y Google OAuth para login de cliente con cuenta Google."),
]
for left, right in rows:
    cells = tech_table.add_row().cells
    set_cell_text(cells[0], left, True)
    set_cell_text(cells[1], right)
    for cell in cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

add_heading(doc, "3. APIs integradas")
doc.add_paragraph("El proyecto integra dos APIs principales:")
add_bullet(doc, "API propia REST: ubicada en la carpeta API. Expone endpoints como /api/products, /api/orders, /api/branches, /api/dashboard y /api/auth. Esta API comunica Angular con MySQL usando JSON.")
add_bullet(doc, "Google OAuth: integrado en el login del cliente. El frontend recibe la credencial de Google y la API la valida en AuthService usando el Client ID configurado en appsettings.json.")

add_heading(doc, "4. CRUD y funcionalidades")
crud_table = doc.add_table(rows=1, cols=3)
crud_table.alignment = WD_TABLE_ALIGNMENT.CENTER
crud_table.style = "Table Grid"
headers = crud_table.rows[0].cells
for i, text in enumerate(["Modulo", "Ubicacion", "Operacion"]):
    set_cell_text(headers[i], text, True)
    shade_cell(headers[i], "F5E6CA")

crud_rows = [
    ("Productos", "Admin: /admin/productos | API: ProductsController", "Crear, leer, editar, eliminar logicamente y subir imagen."),
    ("Pedidos", "Cliente: /cliente/detalle y /cliente/pedidos | API: OrdersController", "Crear pedido, leer historial y actualizar estado desde admin."),
    ("Sucursales", "Admin dashboard | API: BranchesController", "Crear, leer, editar y eliminar sucursales con usuario y contrasena propia."),
    ("Usuarios", "Login y registro | API: AuthController", "Login local, registro de cliente y login con Google."),
]
for row in crud_rows:
    cells = crud_table.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text, i == 0)

add_heading(doc, "5. Seguridad y validaciones")
add_bullet(doc, "Hash de contrasenas: se usa HashHelper con SHA-256 antes de guardar contrasenas en MySQL.")
add_bullet(doc, "Tokens de sesion: TokenHelper genera JWT para mantener la sesion simulada/real del usuario autenticado.")
add_bullet(doc, "Validacion de entradas: los DTOs se validan con FluentValidation en archivos como LoginValidator, ProductoValidator, PedidoValidator y SucursalValidator.")
add_bullet(doc, "Prevencion basica de XSS: SecurityHelper.Clean limpia texto recibido antes de guardar nombres, descripciones y direcciones.")
add_bullet(doc, "Manejo de errores: ApiService centraliza fetch y muestra mensajes controlados en pantalla cuando una solicitud falla.")
add_bullet(doc, "Roles: cliente, admin y sucursal. Las sucursales solo pueden gestionar productos/pedidos, mientras el admin controla sucursales y configuracion.")

add_heading(doc, "6. Arquitectura del proyecto")
add_bullet(doc, "Frontend separado en client/src/app con pages, services y models.")
add_bullet(doc, "API separada en la carpeta API con Controllers, Services, Repositories, Models/DTOs, Models/Entities, Validators, Mappers y Helpers.")
add_bullet(doc, "Fetch centralizado en client/src/app/services/api.service.ts para reutilizar GET, POST, PUT, DELETE y upload.")
add_bullet(doc, "Servicios de Angular por modulo: ProductService, OrderService, BranchService, AuthService y DashboardService.")
add_bullet(doc, "Servicios de API inyectados en controladores: AuthService, ProductosService, PedidosService, SucursalesService y DashboardService.")
add_bullet(doc, "Repositorio generico Repository<T> para reutilizar operaciones de acceso a datos.")
add_bullet(doc, "Estructura modular y escalable, lista para agregar manifest, service worker, cache y notificaciones en una futura PWA.")

add_heading(doc, "7. Flujo de usuario")
add_bullet(doc, "El usuario entra al login y elige si es cliente o negocio.")
add_bullet(doc, "El cliente puede registrarse manualmente o entrar con Google, ver productos, agregar cantidades y confirmar un pedido para recoger en tienda.")
add_bullet(doc, "El administrador entra con correo y contrasena, administra productos, agrega sucursales y revisa pedidos.")
add_bullet(doc, "El pedido inicia como enviado; el administrador puede cambiarlo a en preparacion y despues a listo para recoger.")

add_heading(doc, "8. Relacion con la rubrica")
rubric_table = doc.add_table(rows=1, cols=3)
rubric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rubric_table.style = "Table Grid"
headers = rubric_table.rows[0].cells
for i, text in enumerate(["Criterio", "Cumplimiento", "Evidencia en el proyecto"]):
    set_cell_text(headers[i], text, True)
    shade_cell(headers[i], "F5E6CA")

rubric_rows = [
    ("Desarrollo tecnico 40%", "Angular, asincronia y APIs integradas.", "Componentes Angular, ApiService con fetch, API propia REST y Google OAuth."),
    ("Funcionalidad 20%", "CRUD completo y flujo cliente/admin.", "Productos, pedidos, sucursales, login, dashboard y estados de pedido."),
    ("Seguridad 15%", "Validaciones, proteccion basica y tokens.", "FluentValidation, HashHelper, SecurityHelper y JWT."),
    ("Arquitectura 15%", "Codigo organizado y escalable.", "Separacion client/API, servicios, DTOs, profiles, validators y repositorio generico."),
    ("Presentacion 10%", "Demo y explicacion tecnica.", "Flujo de login, catalogo, pedido, dashboard y administracion."),
]
for row in rubric_rows:
    cells = rubric_table.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text, i == 0)

add_heading(doc, "9. Demo sugerida para presentar")
steps = [
    "Entrar como administrador con admin@sistemaventas.com y contrasena 123456.",
    "Mostrar dashboard con productos activos, pedidos, pendientes y ventas.",
    "Crear una sucursal con nombre, direccion, usuario y contrasena.",
    "Entrar a productos y crear un producto con imagen y precio.",
    "Salir y entrar como cliente registrando una cuenta o usando Google.",
    "Agregar productos al pedido, elegir sucursal y confirmar.",
    "Regresar al admin, abrir pedidos y cambiar el estado a en preparacion y listo para recoger.",
    "Mostrar que los datos persisten al recargar porque se guardan en MySQL.",
]
for i, step in enumerate(steps, start=1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(step)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)

add_heading(doc, "10. Cuenta inicial")
doc.add_paragraph("Administrador inicial: admin@sistemaventas.com")
doc.add_paragraph("Contrasena: 123456")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Proyecto de Tecnologias Web - Plataforma de pedidos")
run.font.name = "Arial"
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(120, 112, 105)

doc.save(OUTPUT)
print(OUTPUT)
